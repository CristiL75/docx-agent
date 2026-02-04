from pathlib import Path
from typing import Dict, Any, TypedDict, List, Optional

from docx import Document
from langgraph.graph import StateGraph, END

from src.docx_io.traverse import iter_text_containers
from src.docx_io.anchors import extract_anchors
from src.docx_io.fill_checkboxes import fill_checkboxes_in_container, fill_checkbox_groups
from src.docx_io.fill_tables import fill_tables, fill_tables_for_anchors
from src.data.normalize import normalize_key, load_json, normalize_data
from src.llm.hf_model import HFModel
from src.llm.map_fields import heuristic_map, composite_map, _infer_field_type, _infer_key_type
from src.validate.mapping_rules import merge_mappings
from src.report.make_report import write_report, write_text_report, build_report
from src.extract_spans import extract_field_spans
from src.map_spans import map_field_spans
from src.fill_docx import fill_spans_in_docx


class State(TypedDict):
    template_path: Path
    data_path: Path
    anchors: List[Dict[str, Any]]
    anchor_clusters: List[Dict[str, Any]]
    field_spans: List[Dict[str, Any]]
    span_mapping: Dict[str, Any]
    suspicious_fills: List[Dict[str, Any]]
    data_norm: Dict[str, Any]
    mapping_heuristic: Dict[str, Dict[str, Any]]
    mapping_llm: Dict[str, Any]
    mapping_final: Dict[str, str]
    mapping_stats: Dict[str, Any]
    issues: List[Dict[str, Any]]
    actions: List[Dict[str, Any]]
    out_path: Path
    report: Dict[str, Any]


class _InternalState(TypedDict, total=False):
    data_raw: Dict[str, Any]
    mapping_heuristic: Dict[str, Dict[str, Any]]
    mapping_llm: Dict[str, Any]


def _normalize_data_node(state: State, artifacts_dir: Path) -> State:
    raw = load_json(state["data_path"])
    data_norm, data_raw = normalize_data(raw)
    state["data_norm"] = data_norm
    state["issues"] = state.get("issues", [])
    internal: _InternalState = {"data_raw": data_raw}
    state.update(internal)
    write_report(artifacts_dir / "data_normalized.json", {"normalized": data_norm, "raw": data_raw})
    return state


def _extract_anchors_node(state: State, artifacts_dir: Path) -> State:
    doc = Document(str(state["template_path"]))
    containers = list(iter_text_containers(doc))
    anchors, clusters = extract_anchors(containers, doc)
    state["anchors"] = anchors
    state["anchor_clusters"] = clusters
    spans = extract_field_spans(containers, doc)
    state["field_spans"] = spans
    anchors_report = []
    for a in anchors:
        cleaned = {k: v for k, v in a.items() if k != "container"}
        options = cleaned.get("options")
        if isinstance(options, list):
            cleaned["options"] = [{kk: vv for kk, vv in opt.items() if kk != "container"} for opt in options]
        anchors_report.append(cleaned)
    write_report(artifacts_dir / "anchors.json", anchors_report)
    write_report(artifacts_dir / "anchor_clusters.json", clusters)
    write_report(artifacts_dir / "field_spans.json", spans)
    return state


def _map_spans_node(state: State, artifacts_dir: Path, model_name: str, seed: int, llm_enabled: bool) -> State:
    model = HFModel(model_name=model_name, seed=seed) if llm_enabled else None
    result = map_field_spans(
        state.get("field_spans", []),
        state.get("data_norm", {}),
        model=model,
        seed=seed,
    )
    state["span_mapping"] = result
    write_report(artifacts_dir / "span_mapping.json", result)
    return state


def _heuristic_map_node(state: State, artifacts_dir: Path, threshold: int) -> State:
    data_keys = list(state["data_norm"].keys())
    heuristic = heuristic_map(state["anchors"], state["data_norm"], data_keys, threshold=0.6)
    state.update({"mapping_heuristic": heuristic})
    write_report(artifacts_dir / "mapping_heuristic.json", heuristic)
    return state


def _llm_map_ambiguous_node(state: State, artifacts_dir: Path, model_name: str) -> State:
    data_keys = list(state["data_norm"].keys())
    model = HFModel(model_name=model_name)
    composite = composite_map(state.get("anchors", []), state["data_norm"], data_keys, model)
    state.update(
        {
            "mapping_heuristic": composite.get("mapping_heuristic", {}),
            "mapping_llm": composite.get("mapping_llm", {}),
            "mapping_final": composite.get("mapping_final", {}),
            "mapping_stats": composite.get("mapping_stats", {}),
        }
    )
    write_report(artifacts_dir / "mapping_heuristic.json", state.get("mapping_heuristic", {}))
    write_report(artifacts_dir / "mapping_llm.json", state.get("mapping_llm", {}))
    write_report(artifacts_dir / "mapping_final.json", state.get("mapping_final", {}))
    return state


def _validate_merge_node(
    state: State,
    artifacts_dir: Path,
    heuristic_threshold: float,
    llm_threshold: float,
    prioritize_llm: bool,
) -> State:
    mapping_final = state.get("mapping_final")
    if not mapping_final:
        mapping_final = merge_mappings(
            state["anchors"],
            state.get("mapping_heuristic", {}),
            state.get("mapping_llm", {}),
            heuristic_threshold=heuristic_threshold,
            llm_threshold=llm_threshold,
            prioritize_llm=prioritize_llm,
        )
        state["mapping_final"] = mapping_final
        write_report(artifacts_dir / "mapping_final.json", mapping_final)

    unmatched = [
        {"label": a.get("label_text"), "location": a.get("location")}
        for a in state["anchors"]
        if a.get("anchor_id") not in mapping_final
    ]
    state["issues"] = unmatched
    return state


def _fill_docx_node(state: State, artifacts_dir: Path, dry_run: bool) -> State:
    if dry_run:
        doc = Document(str(state["template_path"]))
        filled_text, suspicious_fills = fill_spans_in_docx(
            doc,
            state.get("field_spans", []),
            state.get("span_mapping", {}).get("mapping", {}),
            state.get("data_norm", {}),
            state.get("span_mapping", {}).get("computed_values", {}),
            state.get("span_mapping", {}).get("expected_types", {}),
        )
        state["actions"] = state.get("actions", []) + [
            {"type": "summary", "filled_text": filled_text, "filled_checkboxes": 0, "filled_tables": 0}
        ]
        state["suspicious_fills"] = suspicious_fills
        return state

    doc = Document(str(state["template_path"]))
    normalized_data = {normalize_key(k): v for k, v in state["data_norm"].items()}
    mapping = state["mapping_final"]

    actions = state.get("actions", [])

    def _loc_key(location: Dict[str, Any]) -> tuple:
        return (
            location.get("type"),
            location.get("section_idx"),
            location.get("header_footer"),
            location.get("table_idx"),
            location.get("row"),
            location.get("col"),
            location.get("paragraph_idx"),
        )

    location_map = {}
    for container in iter_text_containers(doc):
        location_map[_loc_key(container.location.__dict__)] = container

    checkbox_anchors = [
        a
        for a in state["anchors"]
        if a.get("kind") in {"checkbox", "checkbox_group"} or _infer_field_type(a) == "CHECKBOX_GROUP"
    ]
    table_anchors = [a for a in state["anchors"] if a.get("kind") == "table" or _infer_field_type(a) == "TABLE"]
    text_anchors = [a for a in state["anchors"] if a not in checkbox_anchors and a not in table_anchors]

    filled_text, suspicious_fills = fill_spans_in_docx(
        doc,
        state.get("field_spans", []),
        state.get("span_mapping", {}).get("mapping", {}),
        state.get("data_norm", {}),
        state.get("span_mapping", {}).get("computed_values", {}),
        state.get("span_mapping", {}).get("expected_types", {}),
    )

    label_mapping = {
        a["label_text"]: mapping.get(a["anchor_id"])
        for a in state["anchors"]
        if a.get("label_text") and a.get("anchor_id")
    }

    checkbox_filled = 0
    if checkbox_anchors:
        checkbox_filled += fill_checkbox_groups(checkbox_anchors, state["data_norm"], label_mapping)
        for container in iter_text_containers(doc):
            filled_here = fill_checkboxes_in_container(container, state["data_norm"], label_mapping)
            checkbox_filled += filled_here
            if filled_here:
                actions.append({"type": "checkbox", "label": container.text})

    table_filled = 0
    if table_anchors:
        table_filled = fill_tables_for_anchors(doc, table_anchors, state["data_norm"], mapping)
    elif any(isinstance(v, list) and any(isinstance(i, dict) for i in v) for v in state["data_norm"].values()):
        table_filled = fill_tables(doc, state["data_norm"])

    doc.save(str(state["out_path"]))
    actions.append(
        {
            "type": "summary",
            "filled_text": filled_text,
            "filled_checkboxes": checkbox_filled,
            "filled_tables": table_filled,
        }
    )
    table_json_keys = [
        k
        for k, v in state["data_norm"].items()
        if isinstance(v, list) and any(isinstance(i, dict) for i in v)
    ]
    if checkbox_anchors and checkbox_filled == 0:
        raise SystemExit("CHECKBOX_GROUP anchors present but no checkboxes filled")
    if table_json_keys and table_filled < len(table_json_keys):
        raise SystemExit("TABLE json_keys present but tables not filled")
    for a in state["anchors"]:
        label = str(a.get("label_text") or "")
        nearby = str(a.get("nearby_text") or "")
        norm = f"{label} {nearby}".lower()
        if "catre" not in norm:
            continue
        key = mapping.get(a.get("anchor_id"))
        if not key:
            continue
        key_type = _infer_key_type(key)
        if key_type in {"DATE", "DATE_PARTS"} or key.strip().lower().startswith("data"):
            raise SystemExit("Catre mapped to DATE key")
    state["actions"] = actions
    state["suspicious_fills"] = suspicious_fills
    return state


def _report_node(state: State, artifacts_dir: Path) -> State:
    mapping = state["mapping_final"]
    data_keys = list(state["data_norm"].keys())
    used_keys = {v for v in mapping.values() if v}
    unused_keys = [k for k in data_keys if k not in used_keys]

    llm_candidates = state.get("mapping_llm", {}).get("candidates", {})
    llm_by_anchor = {anchor_id: {"candidates": cands} for anchor_id, cands in llm_candidates.items()}

    mapping_summary = []
    for a in state["anchors"]:
        anchor_id = a.get("anchor_id")
        if not anchor_id:
            continue
        heuristic_item = state.get("mapping_heuristic", {}).get(anchor_id)
        llm_item = llm_by_anchor.get(anchor_id)
        mapping_summary.append(
            {
                "anchor_id": anchor_id,
                "label_text": a.get("label_text"),
                "heuristic": heuristic_item,
                "llm": llm_item,
                "final_key": mapping.get(anchor_id),
            }
        )

    summary = next((a for a in reversed(state.get("actions", [])) if a.get("type") == "summary"), {})
    action_counts = {
        "text": sum(1 for a in state.get("actions", []) if a.get("type") == "text"),
        "checkbox": sum(1 for a in state.get("actions", []) if a.get("type") == "checkbox"),
        "table": int(summary.get("filled_tables", 0)),
    }

    report = build_report(
        anchors_total=len(state["anchors"]),
        filled_text=int(summary.get("filled_text", action_counts["text"])),
        filled_checkboxes=int(summary.get("filled_checkboxes", action_counts["checkbox"])),
        filled_tables=int(summary.get("filled_tables", action_counts["table"])),
        unmatched_anchors=state.get("issues", []),
        unused_json_keys=unused_keys,
        mapping_summary=mapping_summary,
        actions_counts=action_counts,
    )
    span_mapping = state.get("span_mapping", {})
    report["total_field_spans"] = len(state.get("field_spans", []))
    report["unmatched_spans"] = [
        {
            "span_id": s.get("span_id"),
            "left_context": s.get("left_context"),
            "right_context": s.get("right_context"),
            "location": s.get("location"),
        }
        for s in span_mapping.get("unmatched", [])
    ]
    report["type_mismatch_prevented"] = span_mapping.get("type_mismatch_prevented", [])
    report["unmatched_slots"] = report.get("unmatched_spans", [])
    report["suspicious_fills"] = state.get("suspicious_fills", [])
    report["conflicts_found"] = int(state.get("mapping_stats", {}).get("conflicts_found", 0))
    report["repairs_made"] = int(state.get("mapping_stats", {}).get("repairs_made", 0))
    report["role_clusters_detected"] = int(state.get("mapping_stats", {}).get("role_clusters_detected", 0))
    report["role_repairs_made"] = int(state.get("mapping_stats", {}).get("role_repairs_made", 0))

    state["report"] = report
    write_report(artifacts_dir / "report.json", report)
    write_text_report(artifacts_dir / "report.txt", report)
    write_report(artifacts_dir / "actions.json", state.get("actions", []))
    return state


def run_pipeline(
    input_docx: Path,
    input_json: Path,
    output_docx: Path,
    model_name: str,
    artifacts_dir: Path,
    dry_run: bool = False,
    strict: bool = False,
    seed: int = 42,
    llm_enabled: bool = True,
    heuristic_threshold: int = 90,
    llm_threshold: float = 0.15,
    prioritize_llm: bool = True,
    repair_rounds: int = 1,
    repair_heuristic_threshold: int = 80,
    repair_llm_threshold: float = 0.15,
) -> Dict[str, Any]:
    artifacts_dir.mkdir(parents=True, exist_ok=True)

    state: State = {
        "template_path": input_docx,
        "data_path": input_json,
        "anchors": [],
        "anchor_clusters": [],
        "field_spans": [],
        "span_mapping": {},
        "suspicious_fills": [],
        "data_norm": {},
        "mapping_heuristic": {},
        "mapping_llm": {},
        "mapping_final": {},
        "mapping_stats": {},
        "issues": [],
        "actions": [],
        "out_path": output_docx,
        "report": {},
    }

    graph = StateGraph(State)
    graph.add_node("normalize_data_node", lambda s: _normalize_data_node(s, artifacts_dir))
    graph.add_node("extract_anchors_node", lambda s: _extract_anchors_node(s, artifacts_dir))
    graph.add_node("heuristic_map_node", lambda s: _heuristic_map_node(s, artifacts_dir, heuristic_threshold))
    graph.add_node("llm_map_ambiguous_node", lambda s: _llm_map_ambiguous_node(s, artifacts_dir, model_name))
    graph.add_node("map_spans_node", lambda s: _map_spans_node(s, artifacts_dir, model_name, seed, llm_enabled))
    graph.add_node(
        "validate_merge_node",
        lambda s: _validate_merge_node(
            s,
            artifacts_dir,
            heuristic_threshold,
            llm_threshold,
            prioritize_llm,
        ),
    )
    graph.add_node("fill_docx_node", lambda s: _fill_docx_node(s, artifacts_dir, dry_run))
    graph.add_node("report_node", lambda s: _report_node(s, artifacts_dir))

    graph.set_entry_point("normalize_data_node")
    graph.add_edge("normalize_data_node", "extract_anchors_node")
    graph.add_edge("extract_anchors_node", "heuristic_map_node")
    graph.add_edge("heuristic_map_node", "llm_map_ambiguous_node")
    graph.add_edge("llm_map_ambiguous_node", "map_spans_node")
    graph.add_edge("map_spans_node", "validate_merge_node")
    graph.add_edge("validate_merge_node", "fill_docx_node")
    graph.add_edge("fill_docx_node", "report_node")
    graph.add_edge("report_node", END)

    compiled = graph.compile()
    result = compiled.invoke(state)

    for _ in range(max(0, int(repair_rounds))):
        if not result.get("issues"):
            break
        _heuristic_map_node(result, artifacts_dir, threshold=repair_heuristic_threshold)
        _llm_map_ambiguous_node(result, artifacts_dir, model_name)
        _validate_merge_node(
            result,
            artifacts_dir,
            heuristic_threshold=repair_heuristic_threshold,
            llm_threshold=repair_llm_threshold,
            prioritize_llm=prioritize_llm,
        )
        _fill_docx_node(result, artifacts_dir, dry_run)
        _report_node(result, artifacts_dir)

    if strict and result.get("issues"):
        raise SystemExit(2)

    return result
