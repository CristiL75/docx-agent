from typing import Dict, Any, List, Optional, Tuple
import copy

from docx.document import Document
from docx.table import Table, _Row

from src.data.normalize import normalize_key


def _cell_text(cell) -> str:
    return " ".join(cell.text.split())


def _row_is_empty(row: _Row) -> bool:
    return all(not _cell_text(cell) for cell in row.cells)


def _clone_row(table: Table, row: _Row) -> _Row:
    tbl = table._tbl
    new_tr = copy.deepcopy(row._tr)
    tbl.append(new_tr)
    return _Row(new_tr, table)


def _set_cell_text_preserve(cell, value: str) -> None:
    if not cell.paragraphs:
        cell.text = value
        return
    paragraph = cell.paragraphs[0]
    runs = paragraph.runs
    if runs:
        runs[0].text = value
        for run in runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(value)
    for extra in cell.paragraphs[1:]:
        for run in extra.runs:
            run.text = ""


def _header_map(row: _Row) -> Dict[int, str]:
    headers = {}
    for idx, cell in enumerate(row.cells):
        text = _cell_text(cell)
        if text:
            headers[idx] = text
    return headers


def _match_columns(headers: Dict[int, str], sample_row: Dict[str, Any]) -> Dict[int, str]:
    keys = list(sample_row.keys())
    norm_keys = {normalize_key(k): k for k in keys}
    mapping: Dict[int, str] = {}
    for col_idx, header_text in headers.items():
        norm_header = normalize_key(header_text)
        if norm_header in norm_keys:
            mapping[col_idx] = norm_keys[norm_header]
            continue
        for nk, key in norm_keys.items():
            if norm_header in nk or nk in norm_header:
                mapping[col_idx] = key
                break
    return mapping


def _table_has_header(row: _Row, required: List[str]) -> bool:
    text = " ".join(_cell_text(c).lower() for c in row.cells)
    return all(req.lower() in text for req in required)


def _is_services_table(table: Table, prev_text: Optional[str]) -> bool:
    if prev_text and "lista principalelor servicii" in prev_text.lower():
        return True
    if not table.rows:
        return False
    table_text = " ".join(_cell_text(c).lower() for r in table.rows for c in r.cells)
    if "lista principalelor servicii" in table_text:
        return True
    header = table.rows[0]
    return _table_has_header(header, ["nr", "servicii"])


def _is_subcontract_table(table: Table) -> bool:
    if not table.rows:
        return False
    header = table.rows[0]
    return _table_has_header(header, ["nr", "denumire subcontractant"])


def _fill_table(table: Table, rows_data: List[Dict[str, Any]]) -> int:
    if not rows_data or not table.rows:
        return 0
    header_row = table.rows[0]
    headers = _header_map(header_row)
    col_map = _match_columns(headers, rows_data[0])
    if not col_map:
        return 0

    filled_tables = 0
    filled_targets: set = set()
    data_row_idx = 1 if len(table.rows) > 1 else None
    for data_item in rows_data:
        target_row: Optional[_Row] = None
        if data_row_idx is not None and data_row_idx < len(table.rows):
            candidate = table.rows[data_row_idx]
            if _row_is_empty(candidate):
                target_row = candidate
                data_row_idx += 1
        if target_row is None:
            template_row = table.rows[1] if len(table.rows) > 1 else table.rows[0]
            target_row = _clone_row(table, template_row)

        for col_idx, key in col_map.items():
            value = data_item.get(key)
            if value is None:
                continue
            _set_cell_text_preserve(target_row.cells[col_idx], str(value))
            filled += 1

    return filled


def fill_tables(doc: Document, data: Dict[str, Any]) -> int:
    filled = 0
    list_data = {k: v for k, v in data.items() if isinstance(v, list) and any(isinstance(i, dict) for i in v)}

    for table in doc.tables:
        if not table.rows:
            continue
        header_row = table.rows[0]
        headers = _header_map(header_row)
        if not headers:
            continue
        matched = False
        for key, rows in list_data.items():
            if not rows:
                continue
            col_map = _match_columns(headers, rows[0])
            if not col_map:
                continue
            filled += _fill_table(table, rows)
            matched = True
            break
        if matched:
            continue

    return filled


def fill_tables_for_anchors(
    doc: Document,
    anchors: List[Dict[str, Any]],
    data: Dict[str, Any],
    mapping: Dict[str, str],
) -> int:
    filled_tables = 0
    filled_targets: set = set()
    table_cache: Dict[int, Table] = {i: t for i, t in enumerate(doc.tables)}

    for anchor in anchors:
        if str(anchor.get("kind")) != "table":
            continue
        anchor_id = anchor.get("anchor_id")
        if not anchor_id:
            continue
        key = mapping.get(anchor_id)
        if not key:
            continue
        value = data.get(key)
        if not (isinstance(value, list) and any(isinstance(i, dict) for i in value)):
            continue

        loc = anchor.get("location") or {}
        table_idx = loc.get("table_idx")
        target_table = table_cache.get(table_idx) if isinstance(table_idx, int) else None

        if target_table is None:
            for table in doc.tables:
                if not table.rows:
                    continue
                headers = _header_map(table.rows[0])
                if not headers:
                    continue
                if _match_columns(headers, value[0]):
                    target_table = table
                    break

        if target_table is None:
            continue

        if id(target_table) in filled_targets:
            continue
        if _fill_table(target_table, value) > 0:
            filled_targets.add(id(target_table))
            filled_tables += 1

    return filled_tables
