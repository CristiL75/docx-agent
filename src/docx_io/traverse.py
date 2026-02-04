from dataclasses import dataclass
from typing import Iterable, Optional

from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table


@dataclass
class TextContainer:
    kind: str
    obj: object
    text: str
    location: "Location"


@dataclass
class Location:
    type: str
    section_idx: Optional[int] = None
    header_footer: Optional[str] = None
    table_idx: Optional[int] = None
    row: Optional[int] = None
    col: Optional[int] = None
    paragraph_idx: Optional[int] = None


def _iter_paragraphs_in_cell(
    cell: _Cell,
    location_base: Location,
) -> Iterable[TextContainer]:
    for p_idx, paragraph in enumerate(cell.paragraphs):
        location = Location(
            type=location_base.type,
            section_idx=location_base.section_idx,
            header_footer=location_base.header_footer,
            table_idx=location_base.table_idx,
            row=location_base.row,
            col=location_base.col,
            paragraph_idx=p_idx,
        )
        yield TextContainer(kind="table", obj=paragraph, text=paragraph.text, location=location)
    for t_idx, table in enumerate(cell.tables):
        yield from _iter_table_paragraphs(table, location_base, table_idx=t_idx)


def _iter_table_paragraphs(
    table: Table,
    location_base: Location,
    table_idx: int,
) -> Iterable[TextContainer]:
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell_location = Location(
                type=location_base.type,
                section_idx=location_base.section_idx,
                header_footer=location_base.header_footer,
                table_idx=table_idx,
                row=r_idx,
                col=c_idx,
            )
            yield from _iter_paragraphs_in_cell(cell, cell_location)


def _iter_section_paragraphs(doc: Document) -> Iterable[TextContainer]:
    for s_idx, section in enumerate(doc.sections):
        header_base = Location(type="header", section_idx=s_idx, header_footer="header")
        for p_idx, paragraph in enumerate(section.header.paragraphs):
            location = Location(
                type="header",
                section_idx=s_idx,
                header_footer="header",
                paragraph_idx=p_idx,
            )
            yield TextContainer(kind="header_footer", obj=paragraph, text=paragraph.text, location=location)
        for t_idx, table in enumerate(section.header.tables):
            yield from _iter_table_paragraphs(table, header_base, table_idx=t_idx)

        footer_base = Location(type="footer", section_idx=s_idx, header_footer="footer")
        for p_idx, paragraph in enumerate(section.footer.paragraphs):
            location = Location(
                type="footer",
                section_idx=s_idx,
                header_footer="footer",
                paragraph_idx=p_idx,
            )
            yield TextContainer(kind="header_footer", obj=paragraph, text=paragraph.text, location=location)
        for t_idx, table in enumerate(section.footer.tables):
            yield from _iter_table_paragraphs(table, footer_base, table_idx=t_idx)


def iter_blocks(doc: Document) -> Iterable[TextContainer]:
    for p_idx, paragraph in enumerate(doc.paragraphs):
        location = Location(type="body", paragraph_idx=p_idx)
        yield TextContainer(kind="paragraph", obj=paragraph, text=paragraph.text, location=location)
    for t_idx, table in enumerate(doc.tables):
        base = Location(type="body")
        yield from _iter_table_paragraphs(table, base, table_idx=t_idx)
    yield from _iter_section_paragraphs(doc)


def iter_text_containers(doc: Document) -> Iterable[TextContainer]:
    return iter_blocks(doc)
