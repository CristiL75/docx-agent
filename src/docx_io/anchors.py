import re
import string
import unicodedata
from typing import Dict, List, Optional, Tuple

from docx import Document

from src.docx_io.traverse import TextContainer, Location

PLACEHOLDER_UNDERSCORES = re.compile(r"_{4,}")
PLACEHOLDER_DOTS = re.compile(r"\.{4,}")
PLACEHOLDER_DATE = re.compile(r"_{4,}/_{4,}/_{4,}")
CHECKBOX_RE = re.compile(r"(\|_\||\[\s\]|\[x\]|\[X\]|☐|☑)")


def _line_spans(text: str) -> List[Tuple[int, int, str]]:
    spans: List[Tuple[int, int, str]] = []
    cursor = 0
    for line in text.splitlines(keepends=True):
        start = cursor
        end = cursor + len(line)
        spans.append((start, end, line.rstrip("\n")))
        cursor = end
    if not spans:
        spans.append((0, len(text), text))
    return spans


def _find_checkbox_spans(line: str) -> List[Tuple[int, int]]:
    spans: List[Tuple[int, int]] = []
    idx = 0
    while True:
        pos = line.find("|_|", idx)
        if pos == -1:
            break
        spans.append((pos, pos + 3))
        idx = pos + 3
    for match in re.finditer(r"\[\s\]|\[x\]|\[X\]|☐|☑", line):
        spans.append((match.start(), match.end()))
    return spans


def _location_to_dict(location: Location) -> Dict[str, object]:
    return {
        "type": location.type,
        "section_idx": location.section_idx,
        "header_footer": location.header_footer,
        "table_idx": location.table_idx,
        "row": location.row,
        "col": location.col,
        "paragraph_idx": location.paragraph_idx,
    }


def _concat_runs(container: TextContainer) -> Tuple[str, List[Tuple[int, int]]]:
    paragraph = container.obj
    runs = getattr(paragraph, "runs", [])
    spans: List[Tuple[int, int]] = []
    text = ""
    for run in runs:
        start = len(text)
        text += run.text
        end = len(text)
        spans.append((start, end))
    return text, spans


def _find_placeholder_spans(text: str) -> List[Tuple[int, int, str]]:
    spans: List[Tuple[int, int, str]] = []
    occupied = [False] * (len(text) + 1)

    def _reserve(start: int, end: int) -> None:
        for i in range(start, end):
            occupied[i] = True

    for match in PLACEHOLDER_DATE.finditer(text):
        spans.append((match.start(), match.end(), match.group(0)))
        _reserve(match.start(), match.end())

    for regex in (PLACEHOLDER_UNDERSCORES, PLACEHOLDER_DOTS):
        for match in regex.finditer(text):
            if any(occupied[match.start() : match.end()]):
                continue
            spans.append((match.start(), match.end(), match.group(0)))
            _reserve(match.start(), match.end())

    spans.sort(key=lambda s: s[0])
    return spans


def _extract_label_from_text(text: str, start: int) -> str:
    before = text[:start]
    lines = before.splitlines()
    if not lines:
        return ""
    return lines[-1].strip(" :\t")


def _get_table_cell_text(doc: Document, location: Location, col_offset: int) -> Optional[str]:
    if location.table_idx is None or location.row is None or location.col is None:
        return None
    col_idx = location.col + col_offset
    if col_idx < 0:
        return None
    if location.type == "header":
        section = doc.sections[location.section_idx]
        table = section.header.tables[location.table_idx]
    elif location.type == "footer":
        section = doc.sections[location.section_idx]
        table = section.footer.tables[location.table_idx]
    else:
        table = doc.tables[location.table_idx]
    if location.row >= len(table.rows) or col_idx >= len(table.rows[location.row].cells):
        return None
    return table.rows[location.row].cells[col_idx].text.strip()


def extract_anchors(
    containers: List[TextContainer],
    doc: Document,
) -> Tuple[List[Dict[str, object]], List[Dict[str, object]]]:
    anchors: List[Dict[str, object]] = []
    anchor_id = 1
    prev_text: Optional[str] = None
    checkbox_group: List[Dict[str, object]] = []
    checkbox_group_text: List[str] = []
    checkbox_group_location: Optional[Location] = None
    checkbox_gap = 0

    def _flush_checkbox_group() -> None:
        nonlocal anchor_id, checkbox_group, checkbox_group_text, checkbox_group_location
        if not checkbox_group:
            return
        label = "\n".join(checkbox_group_text).strip()
        anchors.append(
            {
                "anchor_id": f"A{anchor_id}",
                "label_text": label,
                "nearby_text": label[:200],
                "placeholder_span": None,
                "location": _location_to_dict(checkbox_group_location) if checkbox_group_location else {},
                "kind": "checkbox_group",
                "options": checkbox_group,
                "container": checkbox_group[0].get("container"),
            }
        )
        anchor_id += 1
        checkbox_group = []
        checkbox_group_text = []
        checkbox_group_location = None

    for container in containers:
        paragraph_text, _ = _concat_runs(container)
        text = paragraph_text
        if not text:
            prev_text = None
            continue

        if CHECKBOX_RE.search(text):
            if not checkbox_group:
                checkbox_group_location = container.location
            checkbox_group_text.append(text.strip())
            checkbox_gap = 0
            for start, end, line in _line_spans(text):
                if not CHECKBOX_RE.search(line):
                    continue
                spans = _find_checkbox_spans(line)
                option_text = (
                    line.replace("|_|", "")
                    .replace("[ ]", "")
                    .replace("[x]", "")
                    .replace("[X]", "")
                    .replace("☐", "")
                    .replace("☑", "")
                    .strip()
                )
                for s, e in spans:
                    checkbox_group.append(
                        {
                            "text": option_text,
                            "span": {"start": start + s, "end": start + e},
                            "container": container,
                        }
                    )
            prev_text = text
            continue
        else:
            if checkbox_group:
                checkbox_gap += 1
                if checkbox_gap > 1:
                    _flush_checkbox_group()
                    checkbox_gap = 0
            else:
                checkbox_gap = 0

        placeholder_spans = _find_placeholder_spans(text)
        for start, end, placeholder in placeholder_spans:
            label = _extract_label_from_text(text, start)

            if container.location.table_idx is not None and not label:
                left_label = _get_table_cell_text(doc, container.location, -1)
                if left_label:
                    label = left_label

            if not label and prev_text:
                if "(" in prev_text and ")" in prev_text:
                    label = prev_text.strip()

            if not label:
                continue

            nearby = text[max(0, start - 100) : min(len(text), end + 100)]
            kind = "table" if container.location.table_idx is not None else "text"

            anchors.append(
                {
                    "anchor_id": f"A{anchor_id}",
                    "label_text": label,
                    "nearby_text": nearby[:200],
                    "placeholder_span": {"start": start, "end": end, "text": placeholder},
                    "location": _location_to_dict(container.location),
                    "kind": kind,
                    "container": container,
                }
            )
            anchor_id += 1

        prev_text = text

    _flush_checkbox_group()

    # cluster anchors by paragraph proximity (window=2)
    clusters: Dict[Tuple, List[Dict[str, object]]] = {}
    for a in anchors:
        loc = a.get("location") or {}
        key = (loc.get("type"), loc.get("section_idx"), loc.get("header_footer"), loc.get("table_idx"))
        clusters.setdefault(key, []).append(a)

    cluster_defs: List[Dict[str, object]] = []
    cluster_id = 1
    for key, items in clusters.items():
        items.sort(key=lambda x: (x.get("location", {}).get("paragraph_idx") or -1))
        current: List[Dict[str, object]] = []
        last_idx = None
        for item in items:
            idx = item.get("location", {}).get("paragraph_idx")
            if last_idx is None or (idx is not None and last_idx is not None and idx - last_idx <= 2):
                current.append(item)
            else:
                cluster_defs.append(_assign_cluster(current, cluster_id))
                cluster_id += 1
                current = [item]
            last_idx = idx
        if current:
            cluster_defs.append(_assign_cluster(current, cluster_id))
            cluster_id += 1

    return anchors, cluster_defs


def _normalize_role_text(value: str) -> str:
    text = value.lower()
    text = "".join(ch for ch in unicodedata.normalize("NFKD", text) if not unicodedata.combining(ch))
    allowed_punct = {"%"}
    table = {
        ch: " "
        for ch in string.punctuation
        if ch not in allowed_punct
    }
    text = text.translate(str.maketrans(table))
    return " ".join(text.split())


def _assign_cluster(items: List[Dict[str, object]], cluster_id: int) -> Dict[str, object]:
    text_context = " ".join(str(i.get("nearby_text") or "") for i in items).strip()
    raw = text_context.lower()
    norm = _normalize_role_text(raw)
    role_pattern = None
    if re.search(r"subsemnatul.*reprezentant.*al", norm, re.DOTALL):
        role_pattern = "PERSON_THEN_ORG"
    elif re.search(r"parafat.*banca.*ziua.*luna.*anul", norm, re.DOTALL):
        role_pattern = "ORG_THEN_DATE"
    elif re.search(r"\(denumirea.*(ofertant|tert|operator)", norm, re.DOTALL):
        role_pattern = "ORG_ONLY"
    elif re.search(r"^catre", norm, re.DOTALL):
        role_pattern = "CATRE_HEADER"
    elif re.search(r"suma.*%|penalitati.*%|reprezentand.*%", norm, re.DOTALL):
        role_pattern = "MONEY_PERCENT"

    for item in items:
        item["cluster_id"] = cluster_id
        item["role_pattern"] = role_pattern

    return {
        "cluster_id": cluster_id,
        "anchors": [i.get("anchor_id") for i in items if i.get("anchor_id")],
        "text_context": text_context,
        "role_pattern": role_pattern,
    }
